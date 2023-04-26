from fastapi import APIRouter, Depends, FastAPI, HTTPException, UploadFile

from app.http.deps import get_db
from app.models.user import User
from app.services.auth import hashing
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import time
import openai
import requests
import json
from PyPDF2 import PdfReader
import docx
import asyncio
import json

router = APIRouter(
    prefix="/gpt"
)


@router.get("/")
async def index():
    return "gpt index"


#  将 YOUR_API_KEY 替换为您的实际 API 密钥
openai.api_key = ""

#  设置API请求的URL和参数
url = "https://api.openai.com/v1/chat/completions"
url2 = "https://api.openai.com/v1/images/generations"

headers = {"authority": "api.openai.com",
           "Content-Type": "application/json",
           "Authorization": "Bearer {}".format(openai.api_key)
           }

proxies = {
    "http": "http://127.0.0.1:7890",
    "https": "http://127.0.0.1:7890",
}


async def ImageCreate(document):
    try:
        # 获取要进行识别的文本内容
        key0 = document.key
        result = imagect(key0)
        print(result)
        results = {"message": "success",
                   "originalText": key0, "correctionResults": result}
        return results
    # 异常处理
    except Exception as e:
        print("异常信息：", e)
        raise HTTPException(status_code=500, detail=str(
            "请求失败，服务器端发生异常！异常信息提示：" + str(e)))


async def TextErrorCorrection(document):
    try:
        # 获取要进行识别的文本内容
        text = document.text
        print(text)
        # 精细分句处理以更好处理长文本
        # data = cut_sent(text) 暂不使用
        key = document.key
        # 进行文本识别和标记
        prompt = 'A：' + text + '//B：' + key
        print(key)
        result = chat(prompt)
        print(result)
        results = {"message": "success",
                   "originalText": document.text, "correctionResults": result}
        return results
    # 异常处理
    except Exception as e:
        print("异常信息：", e)
        raise HTTPException(status_code=500, detail=str(
            "请求失败，服务器端发生异常！异常信息提示：" + str(e)))


async def pdca(document):
    try:
        # 获取要进行识别的文本内容
        value = document.value
        input = document.input
        textarea = document.textarea

        print(textarea)
        # 进行文本识别和标记
        prompt = textarea
        result = aicreate(value, input, prompt)
        print(result)
        results = {"message": "success",
                   "originalText": textarea, "correctionResults": result}
        return results
    # 异常处理
    except Exception as e:
        print("异常信息：", e)
        raise HTTPException(status_code=500, detail=str(
            "请求失败，服务器端发生异常！异常信息提示：" + str(e)))


def get_file_content(filePath):
    with open(filePath, 'rb') as fp:
        return fp.read()


class CommonOcr(object):
    def __init__(self, img_path):
        self._app_id = '0647cbe4b4b53e45ce1386c673207f69'
        self._secret_code = '0bf8e7c1eef74806987e4c5c5fe3615a'
        self._img_path = img_path

    def recognize(self):
        # 通用文字识别
        url = 'https://api.textin.com/ai/service/v2/recognize'
        head = {}
        try:
            image = get_file_content(self._img_path)
            head['x-ti-app-id'] = self._app_id
            head['x-ti-secret-code'] = self._secret_code
            result = requests.post(url, data=image, headers=head)
            return result.text
        except Exception as e:
            return e


async def DocRead(file, key, value):
    # 读取上传的文件
    docBytes = file.file.read()
    docName = file.filename
    print(key)
    array = key.split('，')
    # 判断上传文件类型
    docType = docName.split(".")[-1]
    if docType != "pdf" and docType != "doc" and docType != "docx":
        raise HTTPException(status_code=406, detail=str(
            "请求失败，上传文档格式不正确！请上传pdf或word文档！"))
    try:
        now_time = int(time.mktime(time.localtime(time.time())))
        # 拼接生成随机文件名，注意名称不能包含中文否则后面读取出错
        docPath = "storage/image/" + str(now_time) + "_doc." + docType
        print(docPath)
        doc0 = open(docPath, 'wb')
        doc0.write(docBytes)
        doc0.close()
        txt = ''
        if docType == "pdf":
            reader = PdfReader(docPath)
            # 遍历每一页，并将内容转换为字符串
            page = reader.pages[0]
            txt = page.extract_text()
        else:
            # 打开Word文档
            print('1')
            doc = docx.Document(docPath)
            print('2')
            for para in doc.paragraphs:
                txt += para.text
            print('3')
        print(txt)
        print("文件上传成功！")
        prompt = 'A：' + txt + '//B：' + key
        new_list = chat(prompt)

        # 接口结果返回
        print(new_list)
        results = {"message": "success",
                   "orcResult": "str(ocr_image_results[0])", "correctionResults": new_list}
        return results
    # 异常处理
    except Exception as e:
        print("异常信息：", e)
        raise HTTPException(status_code=500, detail=str(
            "请求失败，服务器端发生异常！异常信息提示：" + str(e)))


async def ImageErrorCorrection(file, key, value):
    # 读取上传的文件
    imgBytes = file.file.read()
    imgName = file.filename
    print(key)
    array = key.split('，')
    # 判断上传文件类型
    imgType = imgName.split(".")[-1]
    if imgType != "png" and imgType != "jpg" and imgType != "jpeg":
        raise HTTPException(status_code=406, detail=str(
            "请求失败，上传图片格式不正确！请上传jpg或png图片！"))
    try:
        now_time = int(time.mktime(time.localtime(time.time())))
        # 拼接生成随机文件名，注意名称不能包含中文否则后面读取出错
        imgPath = "storage/image/" + str(now_time) + "_image." + imgType
        print(imgPath)
        fout = open(imgPath, 'wb')
        fout.write(imgBytes)
        fout.close()
        print("文件上传成功！")
        result = CommonOcr(imgPath)
        data = json.loads(result.recognize())
        lines = data['result']['lines']
        results = ' '.join(line['text'] for line in lines)
        print(results)
        prompt = 'A：' + results + '//B：' + key
        new_list = chat(prompt)

        # 接口结果返回
        print(new_list)
        results1 = {"message": "success",
                    "orcResult": "str(ocr_image_results[0])", "correctionResults": new_list}
        return results1
    # 异常处理
    except Exception as e:
        print("异常信息：", e)
        raise HTTPException(status_code=500, detail=str(
            "请求失败，服务器端发生异常！异常信息提示：" + str(e)))


def chat(prompt):  # 定义一个函数

    try:
        data = {"model": "gpt-3.5-turbo",
                "temperature": 0,
                "messages": [
                    {"role": "system",
                        "content": "你是一个信息提取机器人，接下来我会用：“A：内容//B：信息提示”给你多个prompt，\
                            你需要返回给我一个数组，数组的格式为：\
                                [{\"prompt\":\"prompt的内容\",\"value\":\"提取的内容\"}]"},
                    {"role": "user", "content": "A：甲方购房总价款是12000元，税费为3%//B：甲方购房总价款，甲方购房税费"},
                    {"role": "assistant",
                        "content": "[{\"prompt\":\"甲方购房总价款\",\"value\":\"12000元\"},{\"prompt\":\"甲方购房税费\",\"value\":\"360元\"}]"},
                    {"role": "user", "content": prompt}
                ]
                }
        #  发送HTTP请求
        response = requests.post(url, headers=headers,
                                 json=data, proxies=proxies)
        print(response.text)
        #  解析响应并输出结果
        if response.status_code == 200:
            answer = response.json()[
                "choices"][0]["message"]["content"].strip()
            answer = json.loads(answer)
        else:
            raise Exception(
                f"Request failed with status code {response.text}")
        return answer

    except Exception as exc:
        # print(exc)  #需要打印出故障
        return "broken"


def aicreate(value, input, prompt):  # 定义一个函数

    try:
        data = {}
        if value == 1:
            data = {"model": "gpt-3.5-turbo",
                    "temperature": 1,
                    "messages": [
                        {"role": "system",
                            "content": "你是一个编写日报的机器人，我会告诉你我今天大概的工作内容，\
                                你需要帮我完成markdown格式的日报，格式美观，不需要显示日期，日报的内容需要包括以下项：\
                                    工作成果，工时，经验总结，发现问题，改进措施。每一项的内容包含多条有数字编号的详细说明"},
                        {"role": "user", "content": "项目标题："+input+"，工作概述："+prompt}
                    ]
                    }
        elif value == 2:
            data = {"model": "gpt-3.5-turbo",
                    "temperature": 1,
                    "messages": [
                        {"role": "system",
                         "content": "你是一个编写团建活动报告的机器人，我会告诉你我团建的主题和其他提示信息，\
                             你需要帮我完成markdown格式的团建活动报告的编写，要求格式美观，内容丰富，积极向上\
                                 内容要包括活动日期，参加人员，过程，活动的心得体会等等"},
                        {"role": "user", "content": "团建活动标题："+input+"，活动内容提示："+prompt}
                    ]
                    }
        elif value == 3:
            data = {"model": "gpt-3.5-turbo",
                    "top_p": 0,
                    "messages": [
                        {"role": "system",
                         "content": "你是一个分析体检健康数据的机器人，我会告诉你体检者的一些数据指标，\
                             你需要根据指标帮我完成markdown格式的专业健康分析报告，包括每个指标的解读，\
                                 每项指标检测结果的分析说明，造成结果数值的可能原因以及可能的病症，\
                                     以及针对每项结果的合理化建议"},
                        {"role": "user", "content": prompt}
                    ]
                    }
        #  发送HTTP请求
        response = requests.post(url, headers=headers,
                                 json=data, proxies=proxies)
        print(response.text)
        #  解析响应并输出结果
        if response.status_code == 200:
            answer = response.json()[
                "choices"][0]["message"]["content"].strip()
        else:
            raise Exception(
                f"Request failed with status code {response.text}")
        return answer

    except Exception as exc:
        # print(exc)  #需要打印出故障
        return "broken"


def imagect(prompt):  # 定义一个函数

    try:
        data = {
            "prompt": prompt,
            "n": 1,
            "size": "512x512"
        }

        response = requests.post(
            url2, headers=headers, json=data, proxies=proxies)
        imageurl = response.json()["data"][0]["url"]
        print(imageurl)
        return imageurl
    except Exception as exc:
        # print(exc)  #需要打印出故障
        return "broken"


class Document(BaseModel):
    textarea: str
    value: int
    input: str

# 定义路径操作装饰器：POST方法 + API接口路径

# 文本识别接口


@router.post("/textCorrect/", status_code=200)
# 定义路径操作函数，当接口被访问将调用该函数
async def handle_request(document: Document):
    # 创建一个事件循环
    loop = asyncio.get_running_loop()
    print('创建事件循环成功')
    # 创建一个协程，用于处理当前请求
    coroutine = TextErrorCorrection(document)
    print('创建协程成功')
    # 并行处理当前请求
    result = await asyncio.gather(coroutine, return_exceptions=True)
    print('并行处理当前请求成功')
    print(result)
    return result[0]


# 图片识别接口
@router.post("/imageCorrect/", status_code=200)
# 定义路径操作函数，当接口被访问将调用该函数
async def handle_request(file: UploadFile, key: str, value: int):
    # 创建一个事件循环
    loop = asyncio.get_running_loop()
    print('创建事件循环成功')
    # 创建一个协程，用于处理当前请求
    coroutine = ImageErrorCorrection(file, key, value)
    print('创建协程成功')
    # 并行处理当前请求
    result = await asyncio.gather(coroutine, return_exceptions=True)
    print('并行处理当前请求成功')
    print(result)
    return result[0]

# 文档识别接口


@router.post("/docCorrect/", status_code=200)
# 定义路径操作函数，当接口被访问将调用该函数
async def handle_request(file: UploadFile, key: str, value: int):
    # 创建一个事件循环
    loop = asyncio.get_running_loop()
    print('创建事件循环成功')
    # 创建一个协程，用于处理当前请求
    coroutine = DocRead(file, key, value)
    print('创建协程成功')
    # 并行处理当前请求
    result = await asyncio.gather(coroutine, return_exceptions=True)
    print('并行处理当前请求成功')
    print(result)
    return result[0]


@router.post("/imageCreate/", status_code=200)
# 定义路径操作函数，当接口被访问将调用该函数
async def handle_request(document: Document):
    # 创建一个事件循环
    loop = asyncio.get_running_loop()
    print('创建事件循环成功')
    # 创建一个协程，用于处理当前请求
    coroutine = ImageCreate(document)
    print('创建协程成功')
    # 并行处理当前请求
    result = await asyncio.gather(coroutine, return_exceptions=True)
    print('并行处理当前请求成功')
    print(result)
    return result[0]


@router.post("/pdca/", status_code=200)
# 定义路径操作函数，当接口被访问将调用该函数
async def handle_request(document: Document):
    # 创建一个事件循环
    loop = asyncio.get_running_loop()
    print('创建事件循环成功')
    # 创建一个协程，用于处理当前请求
    coroutine = pdca(document)
    print('创建协程成功')
    # 并行处理当前请求
    result = await asyncio.gather(coroutine, return_exceptions=True)
    print('并行处理当前请求成功')
    print(result)
    return result[0]
